"""Agent 1: Format & Language Detector."""

import logging
from pathlib import Path
from datetime import datetime

import pdfplumber
import docx

import config
from models import FileMetadata

logger = logging.getLogger(__name__)

try:
    from polyglot.detect import Detector
    POLYGLOT_AVAILABLE = True
except ImportError:
    POLYGLOT_AVAILABLE = False
    Detector = None
    logger.warning("polyglot not available, language detection will fail")

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logger.warning("python-magic not available, using file extension detection only")


def detect_file_type(file_path: Path) -> str:
    """Detect file type from file path and content."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix == ".csv":
        return "csv"
    if suffix == ".txt":
        return "txt"
    if suffix == ".json":
        return "json"
    if suffix == ".html":
        return "html"
    if MAGIC_AVAILABLE:
        try:
            mime = magic.Magic(mime=True)
            mime_type = mime.from_file(str(file_path))
            if "pdf" in mime_type:
                return "pdf"
            if "word" in mime_type or "document" in mime_type:
                return "docx"
            if "csv" in mime_type or "text/csv" in mime_type:
                return "csv"
            if "text" in mime_type:
                return "txt"
            if "json" in mime_type:
                return "json"
            if "html" in mime_type:
                return "html"
        except Exception as e:
            logger.warning(f"Could not detect MIME type for {file_path}: {e}")
    return "unknown"


def detect_language(text: str) -> str:
    """Detect language of text content."""
    if not text or len(text.strip()) < 10:
        return "unknown"
    if POLYGLOT_AVAILABLE:
        try:
            detector = Detector(text)
            lang_code = detector.language.code
            if lang_code is None:
                return "unknown"
            return lang_code
        except Exception as e:
            logger.warning(f"Language detection failed: {e}")
    text_lower = text.lower()
    languages = {
        "en": [" the ", " and ", " students ", " course "],
        "fr": [" le ", " la ", " les ", " cours ", " pratique "],
        "de": [" der ", " die ", " und ", " kurs ", " praxis "],
        "lb": [" ze ", " vill ", " praxis ", " cours ", " mir "],
    }
    best_lang = "unknown"
    best_score = 0
    for_lang = list(languages.keys())
    index_lang = 0
    while index_lang < len(for_lang):
        lang = for_lang[index_lang]
        keywords = languages[lang]
        score = 0
        index_kw = 0
        while index_kw < len(keywords):
            keyword = keywords[index_kw]
            if keyword in text_lower:
                score += 1
            index_kw += 1
        if score > best_score:
            best_score = score
            best_lang = lang
        index_lang += 1
    if best_score > 0:
        return best_lang
    return "unknown"


def extract_metadata(file_path: Path) -> dict:
    """Extract file metadata."""
    stat = file_path.stat()
    created_at = datetime.fromtimestamp(stat.st_ctime).isoformat()
    modified_at = datetime.fromtimestamp(stat.st_mtime).isoformat()
    file_size = stat.st_size
    encoding = None
    try:
        with open(file_path, "rb") as f:
            raw_data = f.read(1024)
            if raw_data.startswith(b"\xef\xbb\xbf"):
                encoding = "utf-8-sig"
            elif raw_data.startswith(b"\xff\xfe"):
                encoding = "utf-16-le"
            elif raw_data.startswith(b"\xfe\xff"):
                encoding = "utf-16-be"
            else:
                encoding = "utf-8"
    except Exception as e:
        logger.warning(f"Could not detect encoding for {file_path}: {e}")
        encoding = "unknown"
    return {
        "file_size": file_size,
        "encoding": encoding,
        "created_at": created_at,
        "modified_at": modified_at,
    }


def format_detection_result(file_path: Path, file_type: str, language: str, metadata: dict) -> FileMetadata:
    """Format detection results into FileMetadata object."""
    return FileMetadata(
        file_path=file_path,
        file_type=file_type,
        language=language,
        file_size=metadata["file_size"],
        encoding=metadata.get("encoding"),
        created_at=metadata.get("created_at"),
        modified_at=metadata.get("modified_at"),
    )


def detect_format_and_language(file_path: Path, sample_text: str = None) -> FileMetadata:
    """Main detector agent: detects file format and language."""
    logger.info(f"Detecting format and language for {file_path}")
    file_type = detect_file_type(file_path)
    metadata_dict = extract_metadata(file_path)
    if sample_text is None:
        sample_text = ""
        try:
            encoding = metadata_dict.get("encoding", "utf-8")
            if file_type == "txt" or file_type == "csv":
                with open(file_path, "r", encoding=encoding) as f:
                    sample_text = f.read(1000)
            elif file_type == "pdf":
                with pdfplumber.open(file_path) as pdf:
                    if pdf.pages:
                        page = pdf.pages[0]
                        text_page = page.extract_text() or ""
                        sample_text = text_page[:1000]
            elif file_type == "docx":
                document = docx.Document(file_path)
                para_texts = []
                para_index = 0
                para_limit = min(len(document.paragraphs), 10)
                while para_index < para_limit:
                    para = document.paragraphs[para_index]
                    if para.text:
                        para_texts.append(para.text)
                    para_index += 1
                if para_texts:
                    joined = " ".join(para_texts)
                    sample_text = joined[:1000]
        except Exception as e:
            logger.warning(f"Could not read sample text from {file_path}: {e}")
    language = detect_language(sample_text)
    result = format_detection_result(file_path, file_type, language, metadata_dict)
    logger.info(f"Detected: type={file_type}, language={language}")
    return result

