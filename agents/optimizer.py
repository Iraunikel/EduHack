"""Agent 2: Optimizer - Cleans, deduplicates, and chunks data."""

import logging
import re
import uuid
from pathlib import Path
from typing import Iterable, List
import json
from collections.abc import Mapping, Sequence

import pdfplumber
import docx
from bs4 import BeautifulSoup
from thefuzz import fuzz

import config
from models import ProcessedDocument, TextChunk, FileMetadata

logger = logging.getLogger(__name__)


def extract_text_pdf(file_path: Path) -> str:
    """Extract text from PDF file."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            page_index = 0
            while page_index < len(pdf.pages):
                page = pdf.pages[page_index]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                page_index += 1
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
    return text


def extract_text_docx(file_path: Path) -> str:
    """Extract text from DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        para_index = 0
        while para_index < len(doc.paragraphs):
            para = doc.paragraphs[para_index]
            if para.text:
                text += para.text + "\n"
            para_index += 1
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
    return text


def extract_text_csv(file_path: Path, encoding: str = "utf-8") -> str:
    """Extract text from CSV file."""
    text = ""
    try:
        with open(file_path, "r", encoding=encoding) as f:
            lines = f.readlines()
            line_index = 0
            while line_index < len(lines):
                text += lines[line_index] + "\n"
                line_index += 1
    except Exception as e:
        logger.error(f"Error extracting text from CSV {file_path}: {e}")
    return text


def extract_text_html(file_path: Path, encoding: str = "utf-8") -> str:
    """Extract text from HTML/Notion export file."""
    text = ""
    try:
        with open(file_path, "r", encoding=encoding) as f:
            content = f.read()
        soup = BeautifulSoup(content, "html.parser")
        text = soup.get_text(separator="\n")
    except Exception as e:
        logger.error(f"Error extracting text from HTML {file_path}: {e}")
    return text


def _flatten_json(value) -> Iterable[str]:
    """Yield string representations from arbitrary JSON structures."""
    if value is None:
        return
    if isinstance(value, str):
        yield value
        return
    if isinstance(value, Mapping):
        for item in value.values():
            yield from _flatten_json(item)
        return
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
        for item in value:
            yield from _flatten_json(item)
        return
    yield str(value)


def extract_text(file_path: Path, file_type: str, encoding: str = "utf-8") -> str:
    """Extract text based on file type."""
    if file_type == "pdf":
        return extract_text_pdf(file_path)
    if file_type == "docx":
        return extract_text_docx(file_path)
    if file_type == "csv":
        return extract_text_csv(file_path, encoding)
    if file_type == "html":
        return extract_text_html(file_path, encoding)
    if file_type == "json":
        try:
            with open(file_path, "r", encoding=encoding) as f:
                data = json.load(f)
        except Exception as e:
            logger.error(f"Error parsing JSON file {file_path}: {e}")
            return ""
        return "\n".join(item.strip() for item in _flatten_json(data) if item.strip())
    if file_type == "txt":
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""
    return ""


def clean_text(text: str) -> str:
    """Clean text: remove extra whitespace, encoding issues, boilerplate."""
    if not text:
        return ""
    text = re.sub(r"\x00", "", text)
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = text.strip()
    return text


def detect_duplicates(texts: List[str], threshold: float = None) -> List[int]:
    """Detect duplicate texts using fuzzy matching."""
    if threshold is None:
        threshold = config.FUZZY_MATCH_THRESHOLD
    duplicates = []
    i = 0
    while i < len(texts):
        j = i + 1
        while j < len(texts):
            similarity = fuzz.ratio(texts[i], texts[j]) / 100.0
            if similarity >= threshold:
                duplicates.append(j)
            j += 1
        i += 1
    return duplicates


def chunk_text(text: str, chunk_size: int = None, chunk_overlap: int = None) -> List[str]:
    """Chunk text into smaller pieces."""
    if chunk_size is None:
        chunk_size = config.CHUNK_SIZE
    if chunk_overlap is None:
        chunk_overlap = config.CHUNK_OVERLAP
    if not text:
        return []
    sentences = re.split(r"(?<=[.!?])\s+", text)
    chunks = []
    current_chunk = ""
    sentence_index = 0
    while sentence_index < len(sentences):
        sentence = sentences[sentence_index]
        if len(current_chunk) + len(sentence) <= chunk_size:
            if current_chunk:
                current_chunk += " " + sentence
            else:
                current_chunk = sentence
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence
        sentence_index += 1
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks


def optimize_data(file_path: Path, metadata: FileMetadata) -> ProcessedDocument:
    """Main optimizer agent: extracts, cleans, deduplicates, and chunks data."""
    logger.info(f"Optimizing data from {file_path}")
    raw_text = extract_text(file_path, metadata.file_type, metadata.encoding)
    cleaned_text = clean_text(raw_text)
    lines = cleaned_text.split("\n")
    lines = [line.strip() for line in lines if line.strip()]
    duplicate_indices = detect_duplicates(lines)
    unique_lines = []
    line_index = 0
    while line_index < len(lines):
        if line_index not in duplicate_indices:
            unique_lines.append(lines[line_index])
        line_index += 1
    deduplicated_text = "\n".join(unique_lines)
    chunks_text = chunk_text(deduplicated_text)
    document_id = str(uuid.uuid4())
    chunks = []
    chunk_index = 0
    start_index = 0
    while chunk_index < len(chunks_text):
        chunk_text_item = chunks_text[chunk_index]
        end_index = start_index + len(chunk_text_item)
        chunk_id = f"{document_id}_chunk_{chunk_index}"
        chunk = TextChunk(
            text=chunk_text_item,
            chunk_id=chunk_id,
            document_id=document_id,
            language=metadata.language,
            start_index=start_index,
            end_index=end_index,
            metadata={"file_type": metadata.file_type},
        )
        chunks.append(chunk)
        start_index = end_index
        chunk_index += 1
    result = ProcessedDocument(
        file_path=file_path,
        metadata=metadata,
        raw_text=raw_text,
        cleaned_text=deduplicated_text,
        chunks=chunks,
        duplicates_removed=len(duplicate_indices),
    )
    logger.info(f"Optimized: {len(chunks)} chunks, {len(duplicate_indices)} duplicates removed")
    return result

