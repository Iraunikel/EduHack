"""Generate test data for Edalo pipeline."""

import csv
import logging
import random
from pathlib import Path

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import docx

import config

logger = logging.getLogger(__name__)


def generate_review_templates():
    """Generate review templates with various themes and lengths."""
    templates = {
        "fr": {
            "theory_practice": [
                "Trop de théorie, pas assez de pratique.",
                "Les cours sont intéressants mais manquent d'exercices pratiques.",
                "J'aimerais plus d'activités pratiques et moins de cours magistraux.",
                "Le contenu théorique est bon mais nous avons besoin de plus de pratique.",
                "Les concepts sont bien expliqués mais l'application pratique est limitée.",
                "Il y a un déséquilibre entre la théorie et la pratique dans ce cours.",
                "Nous avons besoin de plus d'exemples concrets et d'exercices appliqués.",
            ],
            "course_structure": [
                "Le rythme du cours est trop rapide.",
                "La structure du cours pourrait être mieux organisée.",
                "Les modules sont bien structurés mais certains sujets manquent de profondeur.",
                "Le programme est complet mais la progression est difficile à suivre.",
                "La séquence des leçons pourrait être améliorée pour une meilleure compréhension.",
            ],
            "materials": [
                "Les supports de cours sont clairs et bien présentés.",
                "Les documents fournis sont utiles mais parfois incomplets.",
                "Nous aurions besoin de plus de ressources complémentaires.",
                "Les matériaux pédagogiques sont de bonne qualité.",
                "Certains supports de cours sont obsolètes et nécessitent une mise à jour.",
            ],
            "engagement": [
                "Le cours est très engageant et interactif.",
                "L'enseignant encourage la participation active des étudiants.",
                "Les discussions en classe sont stimulantes.",
                "Nous aimerions plus d'opportunités d'interaction.",
                "Le niveau d'engagement pourrait être amélioré.",
            ],
        },
        "de": {
            "theory_practice": [
                "Zu viel Theorie, zu wenig Praxis.",
                "Die Kurse sind interessant, aber es fehlen praktische Übungen.",
                "Ich würde gerne mehr praktische Aktivitäten und weniger Vorlesungen haben.",
                "Der theoretische Inhalt ist gut, aber wir brauchen mehr Praxis.",
                "Die Konzepte sind gut erklärt, aber die praktische Anwendung ist begrenzt.",
                "Es gibt ein Ungleichgewicht zwischen Theorie und Praxis in diesem Kurs.",
                "Wir brauchen mehr konkrete Beispiele und angewandte Übungen.",
            ],
            "course_structure": [
                "Das Tempo des Kurses ist zu schnell.",
                "Die Struktur des Kurses könnte besser organisiert sein.",
                "Die Module sind gut strukturiert, aber manche Themen fehlt die Tiefe.",
                "Das Programm ist vollständig, aber der Fortschritt ist schwer zu verfolgen.",
                "Die Reihenfolge der Lektionen könnte verbessert werden für ein besseres Verständnis.",
            ],
            "materials": [
                "Die Kursmaterialien sind klar und gut präsentiert.",
                "Die bereitgestellten Dokumente sind nützlich, aber manchmal unvollständig.",
                "Wir würden mehr ergänzende Ressourcen benötigen.",
                "Die Unterrichtsmaterialien sind von guter Qualität.",
                "Einige Kursmaterialien sind veraltet und müssen aktualisiert werden.",
            ],
            "engagement": [
                "Der Kurs ist sehr ansprechend und interaktiv.",
                "Der Lehrer ermutigt die aktive Teilnahme der Studenten.",
                "Die Klassendiskussionen sind anregend.",
                "Wir würden gerne mehr Möglichkeiten zur Interaktion haben.",
                "Das Engagement-Niveau könnte verbessert werden.",
            ],
        },
        "en": {
            "theory_practice": [
                "Too much theory, not enough practice.",
                "The courses are interesting but lack practical exercises.",
                "I would like more practical activities and fewer lectures.",
                "The theoretical content is good but we need more practice.",
                "The concepts are well explained but practical application is limited.",
                "There is an imbalance between theory and practice in this course.",
                "We need more concrete examples and applied exercises.",
            ],
            "course_structure": [
                "The pace of the course is too fast.",
                "The course structure could be better organized.",
                "The modules are well structured but some topics lack depth.",
                "The program is complete but the progression is difficult to follow.",
                "The sequence of lessons could be improved for better understanding.",
            ],
            "materials": [
                "The course materials are clear and well presented.",
                "The provided documents are useful but sometimes incomplete.",
                "We would need more complementary resources.",
                "The teaching materials are of good quality.",
                "Some course materials are outdated and need updating.",
            ],
            "engagement": [
                "The course is very engaging and interactive.",
                "The teacher encourages active participation from students.",
                "Class discussions are stimulating.",
                "We would like more opportunities for interaction.",
                "The level of engagement could be improved.",
            ],
        },
        "lb": {
            "theory_practice": [
                "Ze vill Theorie, ze wéineg Praxis.",
                "D'Kurse sinn interessant, mee et feelen praktesch Übungen.",
                "Ech géif gär méi praktesch Aktivitéiten a manner Coursen hunn.",
                "De theoreteschen Inhalt ass gutt, mee mir brauchen méi Praxis.",
                "D'Konzepter sinn gutt erkläert, mee d'praktesch Uwendung ass limitéiert.",
                "Et gëtt e Gläichgewiicht tëscht Theorie a Praxis an dësem Cours.",
                "Mir brauchen méi konkreet Beispiller a praktesch Übungen.",
            ],
            "course_structure": [
                "D'Tempo vum Cours ass ze séier.",
                "D'Struktur vum Cours kéint besser organiséiert ginn.",
                "D'Moduler sinn gutt strukturéiert, mee e puer Themen feelen déift.",
                "D'Programm ass komplett, mee d'Progression ass schwéier ze verfollegen.",
                "D'Reiefolleg vun de Léiere kéint verbessert ginn fir e bessert Verständnis.",
            ],
            "materials": [
                "D'Coursmaterialien sinn kloer a gutt presentéiert.",
                "D'Dokumenter déi ginn sinn nëtzlech, mee manchmol onkomplett.",
                "Mir géifen méi ergänzend Ressourcen brauchen.",
                "D'Enseignementsmaterialien sinn vu gudder Qualitéit.",
                "E puer Coursmaterialien sinn veraltet a musse aktualiséiert ginn.",
            ],
            "engagement": [
                "De Cours ass ganz interessant an interaktiv.",
                "Den Enseignant ermëttlecht d'aktiv Participatioun vun de Studenten.",
                "D'Klassendiskussiounen sinn anregend.",
                "Mir géifen gär méi Méiglechkeete fir Interaktioun hunn.",
                "D'Niveau vum Engagement kéint verbessert ginn.",
            ],
        },
    }
    return templates


def generate_reviews(language: str, count: int, templates: dict) -> list:
    """Generate a list of reviews for a given language."""
    reviews = []
    lang_templates = templates.get(language, {})
    themes = list(lang_templates.keys())
    review_index = 0
    duplicate_count = 0
    while review_index < count:
        if review_index < len(themes) * 10:
            theme = themes[review_index % len(themes)]
            theme_reviews = lang_templates.get(theme, [])
            if theme_reviews:
                base_review = theme_reviews[review_index % len(theme_reviews)]
                if review_index % 20 == 0 and duplicate_count < count // 10:
                    reviews.append(base_review)
                    duplicate_count += 1
                reviews.append(base_review)
                review_index += 1
            else:
                review_index += 1
        else:
            theme = random.choice(themes)
            theme_reviews = lang_templates.get(theme, [])
            if theme_reviews:
                reviews.append(random.choice(theme_reviews))
            review_index += 1
    return reviews


def generate_txt_file(output_path: Path, language: str, reviews: list):
    """Generate a text file with multilingual content."""
    logger.info(f"Generating TXT file: {output_path}")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"Student Feedback - Semester 2024 - {language.upper()}\n")
        f.write("=" * 50 + "\n\n")
        review_index = 0
        while review_index < len(reviews):
            f.write(f"Review {review_index + 1}:\n")
            f.write(reviews[review_index] + "\n\n")
            review_index += 1


def generate_csv_file(output_path: Path, language: str, reviews: list):
    """Generate a CSV file with multilingual content."""
    logger.info(f"Generating CSV file: {output_path}")
    with open(output_path, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["student_id", "feedback", "language", "rating", "semester"])
        review_index = 0
        student_id = 1
        while review_index < len(reviews):
            rating = random.randint(2, 5)
            writer.writerow([student_id, reviews[review_index], language, rating, "2024"])
            review_index += 1
            student_id += 1


def generate_docx_file(output_path: Path, language: str, reviews: list):
    """Generate a DOCX file with multilingual content."""
    logger.info(f"Generating DOCX file: {output_path}")
    doc = docx.Document()
    doc.add_heading(f"Student Feedback - {language.upper()} - Semester 2024", 0)
    doc.add_paragraph("Educational Data Collection - Comprehensive Reviews")
    review_index = 0
    while review_index < len(reviews):
        doc.add_heading(f"Review {review_index + 1}", level=2)
        doc.add_paragraph(reviews[review_index])
        review_index += 1
    doc.save(output_path)


def generate_pdf_file(output_path: Path, language: str, reviews: list):
    """Generate a PDF file with multilingual content."""
    logger.info(f"Generating PDF file: {output_path}")
    c = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter
    y_position = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y_position, f"Student Feedback - {language.upper()}")
    y_position -= 30
    c.setFont("Helvetica", 12)
    review_index = 0
    while review_index < len(reviews):
        if y_position < 80:
            c.showPage()
            y_position = height - 50
        c.setFont("Helvetica-Bold", 10)
        c.drawString(50, y_position, f"Review {review_index + 1}:")
        y_position -= 15
        c.setFont("Helvetica", 10)
        lines = reviews[review_index].split("\n")
        line_index = 0
        while line_index < len(lines):
            if y_position < 50:
                c.showPage()
                y_position = height - 50
            c.drawString(60, y_position, lines[line_index][:80])
            y_position -= 12
            line_index += 1
        y_position -= 10
        review_index += 1
    c.save()


def generate_test_data():
    """Generate all test data files with 300+ reviews."""
    logger.info("Generating comprehensive test data with 300+ reviews")
    test_data_dir = config.TEST_DATA_DIR
    test_data_dir.mkdir(parents=True, exist_ok=True)
    templates = generate_review_templates()
    languages = ["fr", "de", "en", "lb"]
    reviews_per_language = 80
    total_reviews = 0
    lang_index = 0
    while lang_index < len(languages):
        language = languages[lang_index]
        reviews = generate_reviews(language, reviews_per_language, templates)
        total_reviews += len(reviews)
        txt_path = test_data_dir / f"feedback_{language}_comprehensive.txt"
        csv_path = test_data_dir / f"feedback_{language}_comprehensive.csv"
        docx_path = test_data_dir / f"feedback_{language}_comprehensive.docx"
        pdf_path = test_data_dir / f"feedback_{language}_comprehensive.pdf"
        generate_txt_file(txt_path, language, reviews)
        generate_csv_file(csv_path, language, reviews)
        generate_docx_file(docx_path, language, reviews)
        generate_pdf_file(pdf_path, language, reviews)
        logger.info(f"Generated {len(reviews)} reviews for {language}")
        lang_index += 1
    logger.info(f"Test data generated in {test_data_dir}")
    logger.info(f"Total reviews generated: {total_reviews}")
    logger.info(f"Total files generated: {len(languages) * 4}")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    generate_test_data()
